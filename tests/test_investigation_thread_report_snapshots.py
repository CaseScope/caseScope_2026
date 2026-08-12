import os
import hashlib
import struct
import tempfile
import zipfile
from types import SimpleNamespace
from unittest.mock import patch

from sqlalchemy import inspect

from config import Config
from models.audit_log import AuditEntityType, AuditLog
from models.database import db
from models.graph import GraphRelationship
from models.graph_saved_view import GraphSavedView
from models.investigation_thread import InvestigationThreadReportSnapshot
from tests.phase0d_test_support import Phase0DSQLiteTestCase, actor, evidence_key
from tests.test_investigation_threads import exact_evidence_payload
from utils.graph_identity import GraphDerivationType, GraphEntityType, GraphRelationshipType
from utils.graph_materializer import clear_case_graph
from utils.graph_saved_views import GraphSavedViewConflictError, GraphSavedViewService
from utils.investigation_thread_report_adapter import InvestigationThreadReportAdapter
from utils.investigation_thread_report_snapshots import GRAPH_HASH_PREFIX, InvestigationThreadReportSnapshotService
from utils.investigation_references import snapshot_sha256
from utils.investigation_threads import InvestigationThreadConflictError, InvestigationThreadService
from utils.report_generator import ReportGenerator
from utils.audit_chain import verify_chain


def png_dimensions(png_bytes):
    assert png_bytes.startswith(b"\x89PNG\r\n\x1a\n")
    width, height = struct.unpack(">II", png_bytes[16:24])
    return width, height


class InvestigationThreadReportSnapshotTestCase(Phase0DSQLiteTestCase):
    def setUp(self):
        super().setUp()
        self.temp_dir = tempfile.TemporaryDirectory()
        self.storage_patch = patch.object(Config, "STORAGE_FOLDER", self.temp_dir.name)
        self.storage_patch.start()
        self.thread_service = InvestigationThreadService()
        self.view_service = GraphSavedViewService()
        self.snapshot_service = InvestigationThreadReportSnapshotService()
        self.actor = actor()

    def tearDown(self):
        self.storage_patch.stop()
        self.temp_dir.cleanup()
        super().tearDown()

    def _thread_with_view(self):
        source, target, relationship = self.make_relationship()
        view = self.view_service.create_view(
            1,
            title="Exact view",
            view_payload={
                "root_entity_ids": [source.id],
                "expanded_entity_ids": [source.id, target.id],
                "visible_relationship_ids": [relationship.id],
                "node_coordinates_json": {
                    str(source.id): {"x": 10, "y": 20},
                    str(target.id): {"x": 300, "y": 20},
                },
            },
            actor=self.actor,
        )["view"]
        thread = self.thread_service.create_thread(1, title="Thread", actor=self.actor)
        version = self.thread_service.link_entity(1, thread["uuid"], expected_version=1, entity_id=source.id, actor=self.actor)["version"]
        version = self.thread_service.link_relationship(1, thread["uuid"], expected_version=version, relationship_id=relationship.id, actor=self.actor)["version"]
        thread = self.thread_service.update_thread(
            1,
            thread["uuid"],
            expected_version=version,
            actor=self.actor,
            current_saved_view_uuid=view["uuid"],
            analyst_conclusion="Analyst conclusion v1",
            include_in_report=True,
        )
        return thread, view

    def test_snapshot_pins_thread_view_fingerprint_hash_and_adapter_is_deterministic(self):
        thread, view = self._thread_with_view()

        snapshot = self.snapshot_service.create_snapshot(
            1,
            thread_uuid=thread["uuid"],
            expected_thread_version=thread["version"],
            expected_saved_view_version=view["version"],
            report_generation_uuid="11111111-1111-4111-8111-111111111111",
            actor=self.actor,
        )
        context = InvestigationThreadReportAdapter().build_context([snapshot])

        self.assertEqual(snapshot["thread_version"], thread["version"])
        self.assertEqual(snapshot["saved_view_version"], view["version"])
        self.assertEqual(snapshot["evidence_set_fingerprint"], thread["evidence_set_fingerprint"])
        self.assertTrue(snapshot["snapshot_sha256"].startswith("report-snapshot:v1:"))
        self.assertTrue(snapshot["graph_render_sha256"].startswith("graph-render:v1:"))
        self.assertTrue(os.path.exists(snapshot["graph_render_path"]))
        self.assertEqual(snapshot["graph_render_format"], "png")
        with open(snapshot["graph_render_path"], "rb") as handle:
            png_bytes = handle.read()
        self.assertEqual(snapshot["graph_render_sha256"], f"{GRAPH_HASH_PREFIX}{hashlib.sha256(png_bytes).hexdigest()}")
        self.assertEqual(png_dimensions(png_bytes), (1200, 800))
        self.assertIn("canonical_graph_svg_sha256", snapshot["snapshot_json"]["render_metadata"])
        self.assertIn("Analyst conclusion v1", context["investigation_thread_report_section"])
        self.assertIn(snapshot["snapshot_sha256"], context["investigation_thread_report_section"])
        self.assertNotIn(snapshot["graph_render_path"], context["investigation_thread_report_section"])

    def test_later_thread_view_and_graph_changes_do_not_change_old_snapshot(self):
        thread, view = self._thread_with_view()
        snapshot = self.snapshot_service.create_snapshot(
            1,
            thread_uuid=thread["uuid"],
            expected_thread_version=thread["version"],
            expected_saved_view_version=view["version"],
            report_generation_uuid="22222222-2222-4222-8222-222222222222",
            actor=self.actor,
        )
        row = InvestigationThreadReportSnapshot.query.filter_by(uuid=snapshot["uuid"]).one()
        original_json = row.snapshot_json
        original_hash = row.snapshot_sha256
        original_graph_hash = row.graph_render_sha256

        updated_thread = self.thread_service.update_thread(
            1,
            thread["uuid"],
            expected_version=thread["version"],
            actor=self.actor,
            analyst_conclusion="Changed v2",
        )
        self.view_service.update_view(
            1,
            view["uuid"],
            expected_version=view["version"],
            actor=self.actor,
            title="Changed view",
        )
        clear_case_graph(1)
        db.session.flush()
        self.make_relationship(case_id=1)
        db.session.commit()

        reloaded = InvestigationThreadReportSnapshot.query.filter_by(uuid=snapshot["uuid"]).one()
        self.assertEqual(reloaded.thread_version, thread["version"])
        self.assertEqual(reloaded.saved_view_version, view["version"])
        self.assertGreater(updated_thread["version"], thread["version"])
        self.assertEqual(reloaded.snapshot_json, original_json)
        self.assertEqual(reloaded.snapshot_sha256, original_hash)
        self.assertEqual(reloaded.graph_render_sha256, original_graph_hash)

    def test_stale_versions_conflict_and_thread_without_saved_view_reports(self):
        source, _target, _relationship = self.make_relationship()
        thread = self.thread_service.create_thread(1, title="No View", actor=self.actor)
        updated = self.thread_service.link_entity(1, thread["uuid"], expected_version=1, entity_id=source.id, actor=self.actor)

        with self.assertRaises(InvestigationThreadConflictError):
            self.snapshot_service.create_snapshot(1, thread_uuid=thread["uuid"], expected_thread_version=1, actor=self.actor)

        snapshot = self.snapshot_service.create_snapshot(
            1,
            thread_uuid=thread["uuid"],
            expected_thread_version=updated["version"],
            actor=self.actor,
        )
        section = InvestigationThreadReportAdapter().build_context([snapshot])["investigation_thread_report_section"]
        self.assertIsNone(snapshot["saved_view_uuid"])
        self.assertIn("No Saved Graph View was captured", section)

    def test_stale_saved_view_version_conflicts(self):
        thread, view = self._thread_with_view()
        self.view_service.update_view(1, view["uuid"], expected_version=view["version"], actor=self.actor, title="v2")

        with self.assertRaises(GraphSavedViewConflictError):
            self.snapshot_service.create_snapshot(
                1,
                thread_uuid=thread["uuid"],
                expected_thread_version=thread["version"],
                expected_saved_view_version=view["version"],
                actor=self.actor,
            )

    def test_docx_without_thread_placeholder_gets_deterministic_thread_fallback(self):
        from docx import Document

        thread, view = self._thread_with_view()
        erk = evidence_key("d")
        with patch(
            "utils.investigation_threads.GraphQueryService.exact_evidence",
            side_effect=lambda case_id, key: exact_evidence_payload(case_id, key),
        ):
            linked = self.thread_service.link_evidence(
                1,
                thread["uuid"],
                expected_version=thread["version"],
                evidence_record_key=erk,
                actor=self.actor,
            )

        snapshot = self.snapshot_service.create_snapshot(
            1,
            thread_uuid=thread["uuid"],
            expected_thread_version=linked["version"],
            expected_saved_view_version=view["version"],
            report_generation_uuid="33333333-3333-4333-8333-333333333333",
            actor=self.actor,
        )
        template_path = os.path.join(self.temp_dir.name, "template.docx")
        output_path = os.path.join(self.temp_dir.name, "output.docx")
        document = Document()
        document.add_paragraph("Minimal legacy report template")
        document.save(template_path)

        ReportGenerator(template_path).generate(InvestigationThreadReportAdapter().build_context([snapshot]), output_path)
        text = "\n".join(paragraph.text for paragraph in Document(output_path).paragraphs)
        current_thread = self.thread_service.get_thread(1, thread["uuid"], include_memberships=False)["thread"]

        self.assertIn("Investigation Thread: Thread", text)
        self.assertIn("Analyst conclusion v1", text)
        self.assertIn(f"Thread version: {linked['version']}", text)
        self.assertIn(f"Saved View version: {view['version']}", text)
        self.assertIn(current_thread["evidence_set_fingerprint"], text)
        self.assertIn(snapshot["snapshot_sha256"], text)
        self.assertIn(erk, text)

    def test_docx_embeds_frozen_saved_graph_and_no_view_thread_still_generates(self):
        from docx import Document

        thread, view = self._thread_with_view()
        snapshot = self.snapshot_service.create_snapshot(
            1,
            thread_uuid=thread["uuid"],
            expected_thread_version=thread["version"],
            expected_saved_view_version=view["version"],
            report_generation_uuid="44444444-4444-4444-8444-444444444444",
            actor=self.actor,
        )
        svg = self.snapshot_service._build_svg(snapshot["snapshot_json"]["saved_view"])
        expected_png = self.snapshot_service._rasterize_svg_to_png(svg)
        expected_graph_hash = f"{GRAPH_HASH_PREFIX}{hashlib.sha256(expected_png).hexdigest()}"
        self.assertIn("ALPHA", svg)
        self.assertIn(r"C:\Tools\cmd.exe", svg)
        self.assertIn("RUNS_IMAGE", svg)
        self.assertIn('marker-end="url(#arrow)"', svg)
        self.assertEqual(snapshot["snapshot_json"]["render_metadata"]["canonical_graph_svg_sha256"], f"{GRAPH_HASH_PREFIX}{snapshot_sha256({'svg': svg})}")
        template_path = os.path.join(self.temp_dir.name, "graph-template.docx")
        output_path = os.path.join(self.temp_dir.name, "graph-output.docx")
        Document().save(template_path)

        ReportGenerator(template_path).generate(InvestigationThreadReportAdapter().build_context([snapshot]), output_path)

        with zipfile.ZipFile(output_path) as package:
            media = [name for name in package.namelist() if name.startswith("word/media/")]
            graph_media = [
                package.read(name)
                for name in media
                if name.lower().endswith(".png") and package.read(name) == expected_png
            ]
        self.assertEqual(len(graph_media), 1)
        self.assertEqual(png_dimensions(graph_media[0]), (1200, 800))
        self.assertGreater(len(graph_media[0]), 1000)
        with open(snapshot["graph_render_path"], "rb") as handle:
            self.assertEqual(graph_media[0], handle.read())
        self.assertEqual(snapshot["graph_render_sha256"], expected_graph_hash)

        no_view_thread = self.thread_service.create_thread(1, title="No View DOCX", actor=self.actor)
        no_view_snapshot = self.snapshot_service.create_snapshot(
            1,
            thread_uuid=no_view_thread["uuid"],
            expected_thread_version=no_view_thread["version"],
            report_generation_uuid="55555555-5555-4555-8555-555555555555",
            actor=self.actor,
        )
        no_view_output = os.path.join(self.temp_dir.name, "no-view-output.docx")
        ReportGenerator(template_path).generate(InvestigationThreadReportAdapter().build_context([no_view_snapshot]), no_view_output)
        self.assertTrue(os.path.exists(no_view_output))

    def test_generated_docx_remains_immutable_after_thread_view_and_graph_changes(self):
        from docx import Document

        thread, view = self._thread_with_view()
        snapshot = self.snapshot_service.create_snapshot(
            1,
            thread_uuid=thread["uuid"],
            expected_thread_version=thread["version"],
            expected_saved_view_version=view["version"],
            report_generation_uuid="12121212-1212-4212-8212-121212121212",
            actor=self.actor,
        )
        template_path = os.path.join(self.temp_dir.name, "immutable-template.docx")
        output_path = os.path.join(self.temp_dir.name, "immutable-output.docx")
        Document().save(template_path)
        ReportGenerator(template_path).generate(InvestigationThreadReportAdapter().build_context([snapshot]), output_path)
        original_docx_hash = hashlib.sha256(open(output_path, "rb").read()).hexdigest()
        original_graph_hash = snapshot["graph_render_sha256"]

        self.thread_service.update_thread(
            1,
            thread["uuid"],
            expected_version=thread["version"],
            actor=self.actor,
            analyst_conclusion="Changed v2",
        )
        self.view_service.update_view(1, view["uuid"], expected_version=view["version"], actor=self.actor, title="Changed view")
        clear_case_graph(1)
        db.session.flush()
        self.make_relationship(case_id=1)
        db.session.commit()

        reopened_text = "\n".join(paragraph.text for paragraph in Document(output_path).paragraphs)
        self.assertEqual(hashlib.sha256(open(output_path, "rb").read()).hexdigest(), original_docx_hash)
        self.assertIn("Analyst conclusion v1", reopened_text)
        self.assertIn(f"Saved View version: {view['version']}", reopened_text)
        self.assertIn(snapshot["evidence_set_fingerprint"], reopened_text)
        self.assertIn(snapshot["snapshot_sha256"], reopened_text)
        self.assertIn(original_graph_hash, reopened_text)
        self.assertNotIn("Changed v2", reopened_text)

    def test_source_availability_is_tri_state(self):
        rows = [
            {"evidence_record_key": evidence_key("a"), "snapshot_available": True},
            {"evidence_record_key": evidence_key("b"), "snapshot_available": True},
            {"evidence_record_key": evidence_key("c"), "snapshot_available": True},
        ]

        def exact(case_id, erk):
            if erk == evidence_key("a"):
                return {"event": {}}
            if erk == evidence_key("b"):
                from utils.graph_query import GraphNotFoundError
                raise GraphNotFoundError("missing")
            raise RuntimeError("clickhouse unavailable")

        with patch("utils.investigation_thread_report_snapshots.GraphQueryService.exact_evidence", side_effect=exact):
            availability = self.snapshot_service._source_availability(1, rows)

        self.assertEqual([item["original_source_status"] for item in availability], ["available", "unavailable", "indeterminate"])
        section = InvestigationThreadReportAdapter._availability_lines(availability)
        self.assertTrue(any("currently retrievable: Yes" in line for line in section))
        self.assertTrue(any("currently retrievable: No" in line for line in section))
        self.assertTrue(any("could not be determined" in line for line in section))

    def test_historical_saved_relationship_renders_after_live_relationship_removed(self):
        thread, view = self._thread_with_view()
        GraphRelationship.query.delete()
        db.session.commit()

        snapshot = self.snapshot_service.create_snapshot(
            1,
            thread_uuid=thread["uuid"],
            expected_thread_version=thread["version"],
            expected_saved_view_version=view["version"],
            report_generation_uuid="66666666-6666-4666-8666-666666666666",
            actor=self.actor,
        )

        frozen = snapshot["snapshot_json"]["saved_view"]["frozen_relationships"][0]
        self.assertFalse(frozen["live_available"])
        self.assertTrue(frozen["source_reference_key"])
        self.assertTrue(frozen["target_reference_key"])
        svg = self.snapshot_service._build_svg(snapshot["snapshot_json"]["saved_view"])
        self.assertIn("RUNS_IMAGE", svg)


class InvestigationThreadReportSnapshotRouteTestCase(Phase0DSQLiteTestCase):
    include_users = True
    register_routes = True
    patch_audit = False

    def setUp(self):
        super().setUp()
        self.temp_dir = tempfile.TemporaryDirectory()
        self.storage_patch = patch.object(Config, "STORAGE_FOLDER", self.temp_dir.name)
        self.storage_patch.start()
        self.thread_service = InvestigationThreadService()
        self.view_service = GraphSavedViewService()
        self.actor = actor()

    def tearDown(self):
        self.storage_patch.stop()
        self.temp_dir.cleanup()
        super().tearDown()

    def _thread_with_snapshot(self):
        source, target, relationship = self.make_relationship()
        view = self.view_service.create_view(
            1,
            title="Route view",
            view_payload={
                "root_entity_ids": [source.id],
                "expanded_entity_ids": [source.id, target.id],
                "visible_relationship_ids": [relationship.id],
            },
            actor=self.actor,
        )["view"]
        thread = self.thread_service.create_thread(1, title="Route Thread", actor=self.actor)
        version = self.thread_service.link_entity(1, thread["uuid"], expected_version=1, entity_id=source.id, actor=self.actor)["version"]
        thread = self.thread_service.update_thread(
            1,
            thread["uuid"],
            expected_version=version,
            actor=self.actor,
            current_saved_view_uuid=view["uuid"],
        )
        snapshot = InvestigationThreadReportSnapshotService().create_snapshot(
            1,
            thread_uuid=thread["uuid"],
            expected_thread_version=thread["version"],
            expected_saved_view_version=view["version"],
            report_generation_uuid="77777777-7777-4777-8777-777777777777",
            actor=self.actor,
            operation_id="88888888-8888-4888-8888-888888888888",
        )
        return thread, view, snapshot

    def test_viewer_can_get_authorized_snapshot_but_cannot_post(self):
        _thread, _view, snapshot = self._thread_with_snapshot()

        self.login_as("viewer_a")
        response = self.client.get(f"/api/investigation/case-a/report-snapshots/{snapshot['uuid']}")
        self.assertEqual(response.status_code, 200)
        self.assertTrue(response.get_json()["success"])

        response = self.client.post(
            "/api/investigation/case-a/report-snapshots",
            json={"thread_uuid": "any", "expected_thread_version": 1},
        )
        self.assertEqual(response.status_code, 403)

    def test_guessed_snapshot_uuid_in_another_case_is_not_found(self):
        thread = self.thread_service.create_thread(2, title="Case B Thread", actor=self.actor)
        snapshot = InvestigationThreadReportSnapshotService().create_snapshot(
            2,
            thread_uuid=thread["uuid"],
            expected_thread_version=thread["version"],
            report_generation_uuid="99999999-9999-4999-8999-999999999999",
            actor=self.actor,
        )

        self.login_as("viewer_a")
        response = self.client.get(f"/api/investigation/case-a/report-snapshots/{snapshot['uuid']}")
        self.assertEqual(response.status_code, 404)

    def test_snapshot_creation_writes_unified_audit_with_operation_id_and_valid_chain(self):
        _thread, _view, snapshot = self._thread_with_snapshot()

        audit = AuditLog.query.filter_by(
            entity_type=AuditEntityType.INVESTIGATION_THREAD_REPORT_SNAPSHOT,
            entity_id=snapshot["uuid"],
        ).one()
        self.assertEqual(audit.operation_id, "88888888-8888-4888-8888-888888888888")
        self.assertTrue(verify_chain()["valid"])
        table_names = set(inspect(db.engine).get_table_names())
        self.assertNotIn("investigation_thread_report_snapshot_audit", table_names)

    def test_snapshot_route_maps_saved_view_conflict_to_409(self):
        thread, view, _snapshot = self._thread_with_snapshot()
        updated = self.view_service.update_view(1, view["uuid"], expected_version=view["version"], actor=self.actor, title="Route view v2")

        self.login_as("analyst")
        response = self.client.post(
            "/api/investigation/case-a/report-snapshots",
            json={
                "thread_uuid": thread["uuid"],
                "expected_thread_version": thread["version"],
                "expected_saved_view_version": view["version"],
            },
        )

        payload = response.get_json()
        self.assertEqual(response.status_code, 409)
        self.assertTrue(payload["stale_version"])
        self.assertEqual(payload["current_version"], updated["view"]["version"])


class DeterministicReportRouteLifecycleTestCase(Phase0DSQLiteTestCase):
    patch_audit = False

    def setUp(self):
        super().setUp()
        self.temp_dir = tempfile.TemporaryDirectory()
        self.storage_patch = patch.object(Config, "STORAGE_FOLDER", self.temp_dir.name)
        self.storage_patch.start()
        self.thread_service = InvestigationThreadService()
        self.view_service = GraphSavedViewService()
        self.actor = actor()

    def tearDown(self):
        self.storage_patch.stop()
        self.temp_dir.cleanup()
        super().tearDown()

    def _snapshot_audits(self, thread_uuid):
        return AuditLog.query.filter(
            AuditLog.entity_type == AuditEntityType.INVESTIGATION_THREAD_REPORT_SNAPSHOT,
            AuditLog.details.like(f'%"thread_uuid": "{thread_uuid}"%'),
        ).order_by(AuditLog.id.asc()).all()

    def _report_thread_with_view(self, title):
        source, target, relationship = self.make_relationship()
        view = self.view_service.create_view(
            1,
            title=f"{title} View",
            view_payload={
                "root_entity_ids": [source.id],
                "expanded_entity_ids": [source.id, target.id],
                "visible_relationship_ids": [relationship.id],
                "node_coordinates_json": {
                    str(source.id): {"x": 120, "y": 140},
                    str(target.id): {"x": 360, "y": 140},
                },
            },
            actor=self.actor,
        )["view"]
        thread = self.thread_service.create_thread(1, title=title, actor=self.actor)
        version = self.thread_service.link_entity(1, thread["uuid"], expected_version=1, entity_id=source.id, actor=self.actor)["version"]
        thread = self.thread_service.update_thread(
            1,
            thread["uuid"],
            expected_version=version,
            actor=self.actor,
            current_saved_view_uuid=view["uuid"],
            include_in_report=True,
        )
        return thread, view

    def _assert_abort_audit(self, thread_uuid, reason):
        audits = self._snapshot_audits(thread_uuid)
        created = [entry for entry in audits if entry.action == "created"]
        deleted = [entry for entry in audits if entry.action == "deleted"]
        self.assertEqual(len(created), 1)
        self.assertEqual(len(deleted), 1)
        self.assertEqual(created[0].operation_id, deleted[0].operation_id)
        self.assertEqual(created[0].to_dict()["details"]["report_generation_uuid"], deleted[0].to_dict()["details"]["report_generation_uuid"])
        self.assertEqual(created[0].to_dict()["details"]["snapshot_uuid"], deleted[0].to_dict()["details"]["snapshot_uuid"])
        self.assertEqual(deleted[0].to_dict()["details"]["reason"], reason)
        self.assertFalse(deleted[0].to_dict()["details"]["completed_report"])
        graph_path = deleted[0].to_dict()["details"].get("graph_render_path")
        if graph_path:
            self.assertFalse(os.path.exists(graph_path))
            root, _ext = os.path.splitext(graph_path)
            self.assertFalse(os.path.exists(f"{root}.svg"))
        self.assertTrue(verify_chain()["valid"])
        table_names = set(inspect(db.engine).get_table_names())
        self.assertNotIn("investigation_thread_report_snapshot_audit", table_names)

    def test_case_report_linkage_failure_cleans_docx_and_snapshots(self):
        from routes import reports as report_routes

        thread, _view = self._report_thread_with_view("Durable Report Thread")
        report_path = os.path.join(self.temp_dir.name, "case-a", "reports", "generated.docx")
        os.makedirs(os.path.dirname(report_path), exist_ok=True)

        def fake_generate_case_report(**_kwargs):
            with open(report_path, "wb") as handle:
                handle.write(b"docx")
            return report_path

        user = SimpleNamespace(
            id=1,
            username="analyst",
            full_name="Analyst",
            permission_level="analyst",
            is_authenticated=True,
            is_administrator=False,
        )
        template = SimpleNamespace(id=10, file_exists=True)

        with (
            self.app.test_request_context("/api/reports/generate/case-a", method="POST", json={"template_id": 10}),
            patch.object(report_routes, "current_user", user),
            patch("models.report_template.ReportTemplate.query") as template_query,
            patch("utils.report_generator.get_base_case_context", return_value={"case_name": "Case A"}),
            patch("utils.hunt_negative_report_adapter.build_negative_findings_report_context", return_value={
                "negative_findings_included": 0,
                "negative_findings_section": "",
                "negative_findings_section_title": "Reviewed Artifacts With No Matching Evidence Identified",
                "negative_findings_audit_appendix": "",
                "negative_findings_audit_appendix_title": "Negative Finding Audit Appendix",
            }),
            patch("utils.report_generator.generate_case_report", side_effect=fake_generate_case_report),
            patch("models.case_report.CaseReport", side_effect=RuntimeError("db failure")),
        ):
            template_query.get.return_value = template
            response, status = report_routes.generate_report.__wrapped__("case-a")

        payload = response.get_json()
        self.assertEqual(status, 500)
        self.assertFalse(payload["success"])
        self.assertFalse(os.path.exists(report_path))
        self.assertEqual(InvestigationThreadReportSnapshot.query.filter_by(thread_uuid=thread["uuid"]).count(), 0)
        self._assert_abort_audit(thread["uuid"], "case_report_registration_failed")

    def test_docx_generation_failure_cleans_created_snapshots(self):
        from routes import reports as report_routes

        thread, _view = self._report_thread_with_view("Render Failure Thread")
        user = SimpleNamespace(
            id=1,
            username="analyst",
            full_name="Analyst",
            permission_level="analyst",
            is_authenticated=True,
            is_administrator=False,
        )
        template = SimpleNamespace(id=10, file_exists=True)

        with (
            self.app.test_request_context("/api/reports/generate/case-a", method="POST", json={"template_id": 10}),
            patch.object(report_routes, "current_user", user),
            patch("models.report_template.ReportTemplate.query") as template_query,
            patch("utils.report_generator.get_base_case_context", return_value={"case_name": "Case A"}),
            patch("utils.hunt_negative_report_adapter.build_negative_findings_report_context", return_value={
                "negative_findings_included": 0,
                "negative_findings_section": "",
                "negative_findings_section_title": "Reviewed Artifacts With No Matching Evidence Identified",
                "negative_findings_audit_appendix": "",
                "negative_findings_audit_appendix_title": "Negative Finding Audit Appendix",
            }),
            patch("utils.report_generator.generate_case_report", return_value=None),
        ):
            template_query.get.return_value = template
            response, status = report_routes.generate_report.__wrapped__("case-a")

        self.assertEqual(status, 500)
        self.assertFalse(response.get_json()["success"])
        self.assertEqual(InvestigationThreadReportSnapshot.query.filter_by(thread_uuid=thread["uuid"]).count(), 0)
        self._assert_abort_audit(thread["uuid"], "docx_generation_failed")

    def test_second_thread_stale_conflict_cleans_first_snapshot(self):
        from routes import reports as report_routes

        first, _view = self._report_thread_with_view("First Thread")
        second = self.thread_service.create_thread(1, title="Second Thread", include_in_report=True, actor=self.actor)
        stale_second_version = second["version"]
        self.thread_service.update_thread(
            1,
            second["uuid"],
            expected_version=second["version"],
            actor=self.actor,
            analyst_conclusion="Updated after caller read",
        )
        user = SimpleNamespace(
            id=1,
            username="analyst",
            full_name="Analyst",
            permission_level="analyst",
            is_authenticated=True,
            is_administrator=False,
        )
        template = SimpleNamespace(id=10, file_exists=True)

        with (
            self.app.test_request_context(
                "/api/reports/generate/case-a",
                method="POST",
                json={
                    "template_id": 10,
                    "thread_versions": {
                        first["uuid"]: first["version"],
                        second["uuid"]: stale_second_version,
                    },
                },
            ),
            patch.object(report_routes, "current_user", user),
            patch("models.report_template.ReportTemplate.query") as template_query,
            patch("utils.report_generator.get_base_case_context", return_value={"case_name": "Case A"}),
            patch("utils.hunt_negative_report_adapter.build_negative_findings_report_context", return_value={
                "negative_findings_included": 0,
                "negative_findings_section": "",
                "negative_findings_section_title": "Reviewed Artifacts With No Matching Evidence Identified",
                "negative_findings_audit_appendix": "",
                "negative_findings_audit_appendix_title": "Negative Finding Audit Appendix",
            }),
        ):
            template_query.get.return_value = template
            response, status = report_routes.generate_report.__wrapped__("case-a")

        self.assertEqual(status, 409)
        self.assertTrue(response.get_json()["stale_version"])
        self.assertEqual(InvestigationThreadReportSnapshot.query.filter_by(thread_uuid=first["uuid"]).count(), 0)
        self._assert_abort_audit(first["uuid"], "stale_thread_version")

    def test_unlicensed_deterministic_report_does_not_invoke_ai_provider(self):
        from routes import reports as report_routes

        thread = self.thread_service.create_thread(1, title="Non-AI Report Thread", actor=self.actor)
        thread = self.thread_service.update_thread(
            1,
            thread["uuid"],
            expected_version=thread["version"],
            actor=self.actor,
            include_in_report=True,
        )
        report_path = os.path.join(self.temp_dir.name, "case-a", "reports", "non-ai.docx")
        os.makedirs(os.path.dirname(report_path), exist_ok=True)

        def fake_generate_case_report(**_kwargs):
            with open(report_path, "wb") as handle:
                handle.write(b"docx")
            return report_path

        user = SimpleNamespace(
            id=1,
            username="analyst",
            full_name="Analyst",
            permission_level="analyst",
            is_authenticated=True,
            is_administrator=False,
        )
        template = SimpleNamespace(id=10, file_exists=True)

        with (
            self.app.test_request_context("/api/reports/generate/case-a", method="POST", json={"template_id": 10}),
            patch.object(report_routes, "current_user", user),
            patch("models.report_template.ReportTemplate.query") as template_query,
            patch("utils.report_generator.get_base_case_context", return_value={"case_name": "Case A"}),
            patch("utils.hunt_negative_report_adapter.build_negative_findings_report_context", return_value={
                "negative_findings_included": 0,
                "negative_findings_section": "",
                "negative_findings_section_title": "Reviewed Artifacts With No Matching Evidence Identified",
                "negative_findings_audit_appendix": "",
                "negative_findings_audit_appendix_title": "Negative Finding Audit Appendix",
            }),
            patch("utils.report_generator.generate_case_report", side_effect=fake_generate_case_report),
            patch("utils.feature_availability.FeatureAvailability.is_ai_enabled", return_value=False),
            patch("utils.ai.router.invoke_text", side_effect=AssertionError("AI provider must not be invoked")),
        ):
            template_query.get.return_value = template
            response = report_routes.generate_report.__wrapped__("case-a")

        payload = response.get_json()
        self.assertTrue(payload["success"])
        snapshot = InvestigationThreadReportSnapshot.query.filter_by(thread_uuid=thread["uuid"]).one()
        self.assertIsNotNone(snapshot.case_report_id)


if __name__ == "__main__":
    import unittest

    unittest.main()
