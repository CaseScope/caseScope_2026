"""Markdown to Word Document Converter

Converts markdown-formatted text from AI responses into proper Word document formatting.

Mapping:
- ## Heading -> Heading 3 style
- ### Heading -> Heading 4 style  
- * item or - item -> Bullet list
- **bold** -> Bold text
- *italic* -> Italic text
- Regular text -> Normal paragraph
"""
import re
from typing import List, Tuple
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docxtpl import DocxTemplate, RichText, Subdoc


def parse_markdown_line(line: str) -> Tuple[str, str, int]:
    """Parse a single line and determine its type.
    
    Returns:
        Tuple of (line_type, content, level)
        - line_type: 'heading', 'bullet', 'numbered', 'paragraph'
        - content: The text content (without markdown syntax)
        - level: For headings (2, 3, 4), for bullets (indent level)
    """
    stripped = line.strip()
    
    if not stripped:
        return ('empty', '', 0)
    
    # Check for headings: ## or ###
    heading_match = re.match(r'^(#{2,4})\s+(.+)$', stripped)
    if heading_match:
        level = len(heading_match.group(1))
        return ('heading', heading_match.group(2), level)
    
    # Check for bullet points: * or -
    bullet_match = re.match(r'^[\*\-•]\s+(.+)$', stripped)
    if bullet_match:
        # Check for indentation level
        indent = len(line) - len(line.lstrip())
        level = indent // 2  # Each 2 spaces = 1 indent level
        return ('bullet', bullet_match.group(1), level)
    
    # Check for numbered list: 1. or 1)
    numbered_match = re.match(r'^\d+[\.\)]\s+(.+)$', stripped)
    if numbered_match:
        return ('numbered', numbered_match.group(1), 0)
    
    # Regular paragraph
    return ('paragraph', stripped, 0)


def apply_inline_formatting(text: str, tpl: DocxTemplate) -> RichText:
    """Apply inline markdown formatting (bold, italic) to text.
    
    Converts **bold** and *italic* to Word formatting.
    
    Args:
        text: The text to format
        tpl: DocxTemplate instance for creating RichText
        
    Returns:
        RichText object with formatting applied
    """
    rt = RichText()
    
    # Pattern to match **bold**, *italic*, or regular text
    # Process in order: bold first (to not confuse with italic)
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))'
    
    pos = 0
    remaining = text
    
    while remaining:
        # Check for bold: **text**
        bold_match = re.match(r'\*\*(.+?)\*\*', remaining)
        if bold_match:
            rt.add(bold_match.group(1), bold=True)
            remaining = remaining[bold_match.end():]
            continue
        
        # Check for italic: *text* (but not **)
        italic_match = re.match(r'\*([^*]+?)\*', remaining)
        if italic_match:
            rt.add(italic_match.group(1), italic=True)
            remaining = remaining[italic_match.end():]
            continue
        
        # Check for inline code: `code`
        code_match = re.match(r'`([^`]+?)`', remaining)
        if code_match:
            rt.add(code_match.group(1), font='Courier New')
            remaining = remaining[code_match.end():]
            continue
        
        # Regular text until next special character
        next_special = re.search(r'[\*`]', remaining)
        if next_special:
            rt.add(remaining[:next_special.start()])
            remaining = remaining[next_special.start():]
        else:
            rt.add(remaining)
            break
    
    return rt


def markdown_to_subdoc(tpl: DocxTemplate, markdown_text: str) -> Subdoc:
    """Convert markdown text to a Word subdocument with proper formatting.
    
    This creates a subdocument that can be inserted into a template placeholder,
    preserving headings, bullet lists, and inline formatting.
    
    Args:
        tpl: DocxTemplate instance
        markdown_text: The markdown-formatted text from AI
        
    Returns:
        Subdoc object to insert into template
    """
    subdoc = tpl.new_subdoc()
    
    lines = markdown_text.split('\n')
    current_list_level = -1  # Track if we're in a list
    
    for line in lines:
        line_type, content, level = parse_markdown_line(line)
        
        if line_type == 'empty':
            # Add empty paragraph for spacing
            subdoc.add_paragraph('')
            current_list_level = -1
            continue
        
        if line_type == 'heading':
            # Map ## -> Heading 3, ### -> Heading 4
            # level 2 (##) -> Heading 3
            # level 3 (###) -> Heading 4
            style_name = f'Heading {level + 1}'
            
            # Apply inline formatting to heading content
            p = subdoc.add_paragraph(style=style_name)
            add_formatted_text(p, content)
            current_list_level = -1
            
        elif line_type == 'bullet':
            # Add as list item
            p = subdoc.add_paragraph(style='List Bullet')
            add_formatted_text(p, content)
            current_list_level = level
            
        elif line_type == 'numbered':
            p = subdoc.add_paragraph(style='List Number')
            add_formatted_text(p, content)
            current_list_level = 0
            
        else:  # paragraph
            p = subdoc.add_paragraph()
            add_formatted_text(p, content)
            current_list_level = -1
    
    return subdoc


def add_formatted_text(paragraph, text: str):
    """Add text with inline formatting to a paragraph.
    
    Handles **bold**, *italic*, and `code` formatting.
    """
    remaining = text
    
    while remaining:
        # Check for bold: **text**
        bold_match = re.match(r'\*\*(.+?)\*\*', remaining)
        if bold_match:
            run = paragraph.add_run(bold_match.group(1))
            run.bold = True
            remaining = remaining[bold_match.end():]
            continue
        
        # Check for italic: *text* (but not **)
        italic_match = re.match(r'\*([^*]+?)\*', remaining)
        if italic_match:
            run = paragraph.add_run(italic_match.group(1))
            run.italic = True
            remaining = remaining[italic_match.end():]
            continue
        
        # Check for inline code: `code`
        code_match = re.match(r'`([^`]+?)`', remaining)
        if code_match:
            run = paragraph.add_run(code_match.group(1))
            run.font.name = 'Courier New'
            remaining = remaining[code_match.end():]
            continue
        
        # Regular text until next special character
        next_special = re.search(r'[\*`]', remaining)
        if next_special:
            paragraph.add_run(remaining[:next_special.start()])
            remaining = remaining[next_special.start():]
        else:
            paragraph.add_run(remaining)
            break


def convert_markdown_sections(tpl: DocxTemplate, sections: dict) -> dict:
    """Convert all markdown sections to subdocuments.
    
    Takes a dictionary of section names to markdown content and returns
    a new dictionary with subdocuments ready for template insertion.
    
    Args:
        tpl: DocxTemplate instance
        sections: Dict of section_name -> markdown_text
        
    Returns:
        Dict of section_name -> Subdoc
    """
    converted = {}
    for name, content in sections.items():
        if content and isinstance(content, str):
            converted[name] = markdown_to_subdoc(tpl, content)
        else:
            converted[name] = content
    return converted
